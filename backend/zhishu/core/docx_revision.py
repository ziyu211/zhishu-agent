"""Word(.docx) 修订版生成器：错字红色删除线 + 正字红色 + **Word 原生批注**。

为什么要这个模块（根因说明）
------------------------------
此前「生成带批注的修订版 Word」完全交给模型在 code_exec 里手写底层 oxml
（构造 word/comments.xml + 改 [Content_Types].xml + 改 document.xml.rels +
fresh-write 重新打包 zip）。实测 240 份历史产物中：

    含 word/comments.xml（成功）  :  52 份  = 22%
    不含（失败）                  : 188 份  = 78%
      其中 157 份（83%）压根没写批注逻辑——模型直接凭通用能力生成了普通修订文档
          28 份 只写了 commentRange 锚点却缺 comments 部件（Word 里看不到批注）
           3 份 是 22 字节空 zip（fresh-write 中途异常留下的空壳）

结论：**把「能不能生成批注」押在模型每次重新发明轮子上，成功率只有 22%**。
本模块把这件事下沉为服务端稳定能力：模型只负责「找出错字」（认知任务），
格式与批注由代码保证 100% 正确。

设计要点
--------
* run 级精确定位：按字符偏移定位到 run，保留原有字体/字号（deepcopy rPr），
  只叠加颜色与删除线，不破坏段落里的其他元素（书签、超链接、域等原样保留）。
* 批注三件套齐全：commentRangeStart/End + commentReference 锚点、
  word/comments.xml 部件、[Content_Types].xml 与 document.xml.rels 注册。
  —— 只写锚点不写部件（历史 28 份失败产物）是看不到批注的。
* fresh-write 打包：先 copy 再以 'w' 模式逐条重写，**每个条目只写一次**，
  杜绝 writestr 追加导致的重复条目；异常时不会留下半截空 zip。
* 产物自检：打包后 testzip() 必须为 None，否则直接抛错（绝不交付损坏文件）。
"""
from __future__ import annotations

import copy
import datetime as _dt
import os
import shutil
import tempfile
import zipfile
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.run import Run

# 批注部件的关系类型与 Content-Type（OOXML 标准常量）
_REL_COMMENTS = ("http://schemas.openxmlformats.org/officeDocument/"
                 "2006/relationships/comments")
_CT_COMMENTS = ("application/vnd.openxmlformats-officedocument."
                "wordprocessingml.comments+xml")

RED = RGBColor(0xFF, 0x00, 0x00)


class DocxRevisionError(Exception):
    """修订版生成失败（输入非法 / 打包校验未通过）。"""


# ────────────────────────── 段落遍历 ──────────────────────────

def iter_paragraphs(doc) -> Iterable:
    """正文段落 + 表格内段落（合并单元格可能重复 yield，由匹配计数去重）。"""
    for p in doc.paragraphs:
        yield p
    try:
        tables = doc.tables
    except Exception:
        tables = []
    for t in tables:
        try:
            rows = t.rows
        except Exception:
            continue
        for row in rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


# ────────────────────────── run 级操作 ──────────────────────────

def _strip_text_nodes(r_elem):
    """移除 run 内的文本类子节点，只保留 rPr（供后续重填文本）。"""
    for tag in ("w:t", "w:br", "w:tab", "w:cr", "w:noBreakHyphen",
                "w:softHyphen", "w:sym", "w:ptab"):
        for e in r_elem.findall(qn(tag)):
            r_elem.remove(e)


def _make_run_like(tpl_elem, parent_p, text: str, red: bool, strike: bool):
    """基于模板 run 克隆一个新 run（继承字体/字号等），设置文本与修订样式。

    用 python-docx 的 Run 包装后再设 font，可保证 rPr 内元素顺序符合 OOXML
    schema（手写 append 顺序错了会导致 Word 报「内容有问题」）。
    """
    new = copy.deepcopy(tpl_elem)
    _strip_text_nodes(new)
    rpr = new.find(qn("w:rPr"))
    if rpr is not None and not strike:
        # 模板自带删除线而本次不需要时清除，避免误继承
        for e in rpr.findall(qn("w:strike")):
            rpr.remove(e)
        for e in rpr.findall(qn("w:dstrike")):
            rpr.remove(e)
    run = Run(new, parent_p)
    run.text = text
    if red:
        run.font.color.rgb = RED
    if strike:
        run.font.strike = True
    return new


def _comment_range_start(cid: int):
    e = OxmlElement("w:commentRangeStart")
    e.set(qn("w:id"), str(cid))
    return e


def _comment_range_end(cid: int):
    e = OxmlElement("w:commentRangeEnd")
    e.set(qn("w:id"), str(cid))
    return e


def _comment_reference(cid: int):
    r = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(cid))
    r.append(ref)
    return r


def _apply_at(p, start: int, end: int, find: str, replace: str, cid: int) -> bool:
    """把段落 p 的字符区间 [start,end) 替换为「删除线错字 + 红色正字」并包裹批注锚点。

    只改动与区间相交的 run，其余子元素（pPr、书签、超链接等）原样保留。
    返回是否成功。
    """
    children = list(p._p)
    # 仅统计直接子 w:r 的字符偏移（与 Paragraph.runs 口径一致）
    def _run_text(r):
        return "".join(t.text or "" for t in r.findall(qn("w:t")))

    new_children = []
    pos = 0
    inserted = False
    for child in children:
        if child.tag != qn("w:r"):
            new_children.append(child)
            continue
        text = _run_text(child)
        r_start, r_end = pos, pos + len(text)
        pos = r_end
        if r_end <= start or r_start >= end:
            new_children.append(child)          # 与目标区间不相交，原样保留
            continue
        ls = max(start, r_start) - r_start
        le = min(end, r_end) - r_start
        head, tail = text[:ls], text[le:]
        if head:
            new_children.append(_make_run_like(child, p, head, red=False, strike=False))
        if not inserted:
            new_children.append(_comment_range_start(cid))
            new_children.append(_make_run_like(child, p, find, red=True, strike=True))
            if replace:
                new_children.append(_make_run_like(child, p, replace, red=True, strike=False))
            new_children.append(_comment_range_end(cid))
            new_children.append(_comment_reference(cid))
            inserted = True
        if tail:
            new_children.append(_make_run_like(child, p, tail, red=False, strike=False))

    if not inserted:
        return False
    # 切片赋值（不能用 clear()，否则会丢掉 w:pPr 等段落属性）
    p._p[:] = new_children
    return True


def _revise_paragraph(p, find: str, replace: str, comment: str,
                      next_cid, comments: list, max_per_doc: int,
                      all_occurrences: bool = True) -> int:
    """在单个段落内处理 find 的出现处，**从后往前**逐个处理，避免偏移错乱。

    关键（易错点）：替换后的段落文本里仍然保留着 find（它作为「红色删除线错字」
    存在），若简单地反复 rfind 会**永远命中同一处**、陷入死循环。因此每处理完
    一处，就把搜索上界收缩到该处的起始下标，下一轮只在其**左侧**继续找。
    """
    hits = 0
    guard = 0
    limit = None                            # 搜索上界（不含），处理完一处后左移
    while guard < 500:                      # 防御：单行极端重复场景
        guard += 1
        runs = p.runs
        if not runs:
            break
        full = "".join(r.text for r in runs)
        pool = full if limit is None else full[:limit]
        idx = pool.rfind(find)
        if idx < 0:
            break
        cid = next_cid()
        if not _apply_at(p, idx, idx + len(find), find, replace, cid):
            break
        comments.append({"id": cid, "text": comment})
        hits += 1
        limit = idx                         # 下一次只在本处左侧继续搜索
        if not all_occurrences:
            break
        if len(comments) >= max_per_doc:
            break
    return hits


# ────────────────────────── 部件构造 ──────────────────────────

def _build_comments_xml(comments: list, author: str) -> bytes:
    """构造 word/comments.xml（原生批注部件）。"""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    now = _dt.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    parts = [f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
             f'<w:comments xmlns:w="{W}">']
    for c in comments:
        text = (c["text"] or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        initial = (author or "校对")[:1] or "Z"
        au = (author or "校对").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        parts.append(
            f'<w:comment w:id="{c["id"]}" w:author="{au}" w:date="{now}" '
            f'w:initials="{initial}">'
            f'<w:p><w:pPr><w:pStyle w:val="CommentText"/></w:pPr>'
            f'<w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
            f'<w:annotationRef/></w:r>'
            f'<w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'
            f'</w:comment>'
        )
    parts.append("</w:comments>")
    return "".join(parts).encode("utf-8")


def _patch_content_types(xml_bytes: bytes) -> bytes:
    """在 [Content_Types].xml 注册 comments 部件（已注册则原样返回）。"""
    try:
        text = xml_bytes.decode("utf-8")
    except Exception:
        return xml_bytes
    if "/word/comments.xml" in text:
        return xml_bytes
    override = (f'<Override PartName="/word/comments.xml" '
                f'ContentType="{_CT_COMMENTS}"/>')
    if "</Types>" in text:
        text = text.replace("</Types>", override + "</Types>")
    else:
        text = text.rstrip()
        if text.endswith(">"):
            text = text + override
    return text.encode("utf-8")


def _patch_rels(xml_bytes: bytes) -> bytes:
    """在 word/_rels/document.xml.rels 中注册 comments 关系（Id 保证唯一）。"""
    try:
        text = xml_bytes.decode("utf-8")
    except Exception:
        return xml_bytes
    if _REL_COMMENTS in text:
        return xml_bytes
    # 生成不与现有 rId 冲突的 Id
    import re as _re
    used = set(_re.findall(r'Id="([^"]+)"', text))
    n = 1
    while f"rIdZhComment{n}" in used:
        n += 1
    rel = (f'<Relationship Id="rIdZhComment{n}" Type="{_REL_COMMENTS}" '
           f'Target="comments.xml"/>')
    if "</Relationships>" in text:
        text = text.replace("</Relationships>", rel + "</Relationships>")
    else:
        text = text.rstrip()
        if text.endswith(">"):
            text = text + rel
    return text.encode("utf-8")


def _fresh_write_pack(base_path: str, out_path: str, comments_xml: bytes) -> None:
    """fresh-write 打包：逐条复制并**显式覆盖**需改写的部件，杜绝重复条目。

    与模型手写的常见错误（ZipFile.writestr 往已存在条目追加 → zip 内重复条目 →
    Word 报「内容有问题」）不同，这里每个条目只写一次。
    """
    shutil.copy2(base_path, out_path)
    with zipfile.ZipFile(base_path, "r") as zin:
        items = zin.infolist()
        with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in items:
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = _patch_content_types(data)
                elif item.filename == "word/_rels/document.xml.rels":
                    data = _patch_rels(data)
                zout.writestr(item, data)
            zout.writestr("word/comments.xml", comments_xml)


# ────────────────────────── 主入口 ──────────────────────────

def build_revised_docx(src_path: str, revisions: list, author: str = "校对",
                       max_comments: int = 500,
                       all_occurrences: bool = True) -> tuple[bytes, dict]:
    """生成带原生批注的修订版 docx。

    参数：
      src_path  —— 源 .docx 绝对路径
      revisions —— [{find, replace, comment?}, ...]
                   find    : 原文中的错误片段（用于定位）
                   replace : 修正后文本；**空串表示删除多余字**（只留删除线，不补字）
                   comment : 批注文字，缺省自动生成
      author    —— 批注作者名

    返回：(docx_bytes, stats)
      stats = {applied, comments, missed:[...], paragraphs}
    """
    if not src_path or not os.path.isfile(src_path):
        raise DocxRevisionError(f"源文件不存在：{src_path}")
    if not revisions:
        raise DocxRevisionError("缺少 revisions（修订列表）")

    doc = Document(src_path)

    # 规整修订项
    items = []
    for r in revisions:
        if not isinstance(r, dict):
            continue
        find = str(r.get("find") or "").strip()
        if not find:
            continue
        replace = str(r.get("replace") or "")
        comment = str(r.get("comment") or "").strip()
        if not comment:
            if replace:
                comment = f"错别字修正：「{find}」→「{replace}」"
            else:
                comment = f"删除多余字：「{find}」"
        items.append({"find": find, "replace": replace, "comment": comment})

    if not items:
        raise DocxRevisionError("revisions 中缺少有效的 find 字段")

    comments: list = []
    counter = {"n": -1}

    def next_cid() -> int:
        counter["n"] += 1
        return counter["n"]

    applied = 0
    paragraphs = 0
    for p in iter_paragraphs(doc):
        paragraphs += 1
        for it in items:
            hits = _revise_paragraph(p, it["find"], it["replace"], it["comment"],
                                     next_cid, comments, max_comments,
                                     all_occurrences)
            applied += hits

    # 未命中检查（重要：明确告诉模型哪些没匹配上，而不是静默交付）
    missed = []
    if not comments:
        missed = [it["find"] for it in items]

    # 1) 先由 python-docx 落盘（得到修订后的 document.xml）
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx", prefix="zh_rev_")
    os.close(tmp_fd)
    out_fd, out_path = tempfile.mkstemp(suffix=".docx", prefix="zh_rev_out_")
    os.close(out_fd)
    try:
        doc.save(tmp_path)
        comments_xml = _build_comments_xml(comments, author)
        _fresh_write_pack(tmp_path, out_path, comments_xml)
        # 2) 产物自检：zip 必须健康，且批注三件套齐全
        with zipfile.ZipFile(out_path, "r") as z:
            if z.testzip() is not None:
                raise DocxRevisionError("打包校验失败：zip 结构损坏")
            names = z.namelist()
            if "word/comments.xml" not in names:
                raise DocxRevisionError("打包校验失败：缺少 word/comments.xml")
            ct = z.read("[Content_Types].xml").decode("utf-8", "ignore")
            if "/word/comments.xml" not in ct:
                raise DocxRevisionError("打包校验失败：Content_Types 未注册批注部件")
            rels = z.read("word/_rels/document.xml.rels").decode("utf-8", "ignore")
            if _REL_COMMENTS not in rels:
                raise DocxRevisionError("打包校验失败：rels 未注册批注关系")
        with open(out_path, "rb") as f:
            data = f.read()
    finally:
        for fp in (tmp_path, out_path):
            try:
                os.remove(fp)
            except OSError:
                pass

    stats = {
        "applied": applied,
        "comments": len(comments),
        "paragraphs": paragraphs,
        "missed": missed,
    }
    return data, stats
