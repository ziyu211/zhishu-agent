"""智枢智能体 —— 知识库 RAG 模块。

负责：文档切分 → 离线 Embedding → 本地向量库写入；查询时检索 top-k 作为上下文。
纯内网离线，不依赖任何外部检索服务。

本次完善：
  * 文档级元数据（标题/来源/类型/归属人/分块数/字符数/大小/预览正文）登记，
    支撑「文档列表 / 预览 / 删除」前端展示（修复上传后无法显示的问题）。
  * 上传支持更多格式：TXT/MD/CSV/JSON/代码/日志等文本类，以及可选的
    PDF(.pdf) / Word(.docx) / Excel(.xlsx)（需安装对应解析库，缺失则友好报错）。
  * 归属隔离：owner 为空表示共享文档（对所有用户可见/可检索），非空表示私有，
    仅归属人及管理员可见、可检索、可删除。
"""
from __future__ import annotations

import os
import re
import io
import uuid
import zipfile
import shutil
import subprocess
import tempfile
import threading
import xml.etree.ElementTree as ET
from typing import List, Optional, Tuple

from .embedding import EmbeddingEngine
from .vector_store import VectorStore
from .kgraph import KnowledgeGraph
from .config import EmbeddingConfig, VectorStoreConfig


# ── 标准库（零依赖）Office 文档提取 ───────────────────────────────────────
# 借鉴 hermes-agent read_extract 的设计：优先用 zipfile + xml 直接解析
# .docx/.xlsx，不依赖 python-docx / openpyxl，避免「缺库即要求装插件」与
# 潜在的环境耦合。仅在标准库提取为空时才回退到第三方库。

def _local(tag: str) -> str:
    """去掉 XML 命名空间前缀，取本地标签名。"""
    return tag.split("}", 1)[-1] if "}" in tag else tag


def _elem_text(elem) -> str:
    """收集元素及其所有后代中 <t>/<si> 内的纯文本（按文档顺序拼接）。"""
    parts = []
    for t in elem.iter():
        if _local(t.tag) in ("t", "si"):
            parts.append(t.text or "")
    return "".join(parts)


def _extract_docx_stdlib(raw: bytes) -> str:
    """用标准库从 .docx 提取正文与表格文本（零依赖）。失败返回空串。"""
    try:
        z = zipfile.ZipFile(io.BytesIO(raw))
    except Exception:
        return ""
    names = set(z.namelist())
    main = "word/document.xml" if "word/document.xml" in names else None
    if not main:
        for n in names:
            if n.endswith("document.xml"):
                main = n
                break
    if not main:
        return ""
    targets = [main]
    for n in names:
        if (n.startswith("word/header") or n.startswith("word/footer")) and n.endswith(".xml"):
            targets.append(n)

    blocks: List[str] = []
    for part in targets:
        try:
            root = ET.fromstring(z.read(part))
        except Exception:
            continue
        # 找到 body（无则直接用根）
        body = None
        for el in root.iter():
            if _local(el.tag) == "body":
                body = el
                break
        body = body or root

        def walk(elem):
            for child in elem:
                ln = _local(child.tag)
                if ln == "p":
                    line = _elem_text(child).strip()
                    if line:
                        blocks.append(line)
                elif ln == "tbl":
                    rows = []
                    for tr in child:
                        if _local(tr.tag) != "tr":
                            continue
                        cells = []
                        for tc in tr:
                            if _local(tc.tag) != "tc":
                                continue
                            cells.append(_elem_text(tc).strip())
                        if any(cells):
                            rows.append(" | ".join(cells))
                    if rows:
                        blocks.append("\n".join(rows))
                else:
                    walk(child)

        walk(body)
    return "\n".join(blocks).strip()


def _extract_xlsx_stdlib(raw: bytes) -> str:
    """用标准库从 .xlsx 提取全部工作表文本（零依赖）。失败返回空串。"""
    try:
        z = zipfile.ZipFile(io.BytesIO(raw))
    except Exception:
        return ""
    names = set(z.namelist())

    # 1) 共享字符串表
    shared: List[str] = []
    if "xl/sharedStrings.xml" in names:
        try:
            sroot = ET.fromstring(z.read("xl/sharedStrings.xml"))
            for si in sroot:
                if _local(si.tag) != "si":
                    continue
                shared.append(_elem_text(si).strip())
        except Exception:
            pass

    # 2) 工作簿顺序与表文件映射
    R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    sheet_order: List[tuple] = []
    if "xl/workbook.xml" in names:
        try:
            wb = ET.fromstring(z.read("xl/workbook.xml"))
            rels: dict = {}
            if "xl/_rels/workbook.xml.rels" in names:
                rroot = ET.fromstring(z.read("xl/_rels/workbook.xml.rels"))
                for rel in rroot:
                    rels[rel.get("Id")] = rel.get("Target") or ""
            for sheets in wb.iter():
                if _local(sheets.tag) != "sheets":
                    continue
                for sh in sheets:
                    if _local(sh.tag) != "sheet":
                        continue
                    name = sh.get("name") or "Sheet"
                    rid = sh.get(f"{{{R_NS}}}id")
                    target = (rels.get(rid) or "").lstrip("/")
                    if not target.startswith("xl/"):
                        target = "xl/" + target
                    sheet_order.append((name, target))
        except Exception:
            pass
    if not sheet_order:
        for n in names:
            if n.startswith("xl/worksheets/") and n.endswith(".xml"):
                sheet_order.append((n, n))

    # 3) 逐表提取
    out: List[str] = []
    for name, sheetfile in sheet_order:
        try:
            sroot = ET.fromstring(z.read(sheetfile))
        except Exception:
            continue
        rows = []
        for row in sroot.iter():
            if _local(row.tag) != "row":
                continue
            cells = []
            for c in row:
                if _local(c.tag) != "c":
                    continue
                t = c.get("t")
                v = None
                for child in c:
                    ln = _local(child.tag)
                    if ln == "v":
                        v = child.text
                    elif ln == "is":
                        v = _elem_text(child).strip()
                if v is None:
                    cells.append("")
                elif t == "s":
                    try:
                        cells.append(shared[int(v)])
                    except (ValueError, IndexError):
                        cells.append(v or "")
                else:
                    cells.append(v or "")
            if any(str(x).strip() for x in cells):
                rows.append("\t".join(str(x) for x in cells))
        if rows:
            out.append(f"### {name}")
            out.append("\n".join(rows))
    return "\n\n".join(out).strip()


def _extract_pptx_stdlib(raw: bytes) -> str:
    """用标准库从 .pptx 提取全部幻灯片文本（零依赖）。失败返回空串。"""
    try:
        z = zipfile.ZipFile(io.BytesIO(raw))
    except Exception:
        return ""
    names = sorted(z.namelist())
    slides = [n for n in names if re.match(r"ppt/slides/slide\d+\.xml$", n)]
    slides.sort(key=lambda n: int(re.search(r"(\d+)", n).group(1)))
    if not slides:
        slides = [n for n in names if n.startswith("ppt/slides/") and n.endswith(".xml")]
    blocks: List[str] = []
    for idx, s in enumerate(slides, 1):
        try:
            root = ET.fromstring(z.read(s))
        except Exception:
            continue
        texts = [t.text or "" for t in root.iter() if _local(t.tag) == "t"]
        line = " ".join(x.strip() for x in texts if x.strip())
        if line:
            blocks.append(f"[幻灯片 {idx}]\n{line}")
    return "\n\n".join(blocks).strip()


def _extract_odf_stdlib(raw: bytes) -> str:
    """用标准库从 OpenDocument(.odt/.ods/.odp) 提取文本（零依赖）。失败返回空串。"""
    try:
        z = zipfile.ZipFile(io.BytesIO(raw))
    except Exception:
        return ""
    if "content.xml" not in z.namelist():
        return ""
    try:
        root = ET.fromstring(z.read("content.xml"))
    except Exception:
        return ""
    blocks: List[str] = []
    for p in root.iter():
        if _local(p.tag) == "p":
            txt = "".join(p.itertext()).strip()
            if txt:
                blocks.append(txt)
    for tbl in root.iter():
        if _local(tbl.tag) == "table":
            rows = []
            for row in tbl:
                if _local(row.tag) != "table-row":
                    continue
                cells = []
                for cell in row:
                    if _local(cell.tag) != "table-cell":
                        continue
                    cells.append("".join(cell.itertext()).strip())
                if any(cells):
                    rows.append("\t".join(cells))
            if rows:
                blocks.append("\n".join(rows))
    return "\n".join(blocks).strip()


def _extract_rtf(raw: bytes) -> str:
    """尽力从 .rtf 提取纯文本（零依赖，控制字剥离）。失败返回空串。"""
    try:
        text = raw.decode("latin-1")
    except Exception:
        return ""
    text = text.replace("\r", " ").replace("\n", " ")
    # 丢弃常见二进制/样式目标组，避免大量乱码
    for dest in ("fonttbl", "colortbl", "stylesheet", "info", "pict", "shppict", "generator"):
        text = re.sub(r"\{\\*?\\?" + dest + r"[^{}]*\}", " ", text)
    # 字符转义 \'xx
    def _hex(m):
        try:
            return bytes([int(m.group(1), 16)]).decode("latin-1")
        except Exception:
            return ""
    text = re.sub(r"\\'([0-9a-fA-F]{2})", _hex, text)
    # 常见控制符号 → 文本
    text = (text.replace("\\par", "\n").replace("\\line", "\n")
                .replace("\\tab", "\t").replace("\\bullet", "•")
                .replace("\\ldblquote", '"').replace("\\rdblquote", '"')
                .replace("\\lquote", "'").replace("\\rquote", "'")
                .replace("\\emdash", "—").replace("\\endash", "–"))
    # 剩余控制字（\word 及可选数值）
    text = re.sub(r"\\[a-zA-Z]+\-?\d* ?", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\\[^a-zA-Z]", "", text)
    lines = [ln.strip() for ln in text.split("\n")]
    return "\n".join(ln for ln in lines if ln).strip()


def _extract_epub_stdlib(raw: bytes) -> str:
    """用标准库从 .epub 提取正文文本（零依赖）。失败返回空串。"""
    import posixpath
    try:
        z = zipfile.ZipFile(io.BytesIO(raw))
    except Exception:
        return ""
    names = set(z.namelist())
    opf_path = None
    try:
        croot = ET.fromstring(z.read("META-INF/container.xml"))
        for rf in croot.iter():
            if _local(rf.tag) == "rootfile":
                opf_path = rf.get("full-path")
                break
    except Exception:
        pass
    if not opf_path:
        for n in names:
            if n.endswith(".opf"):
                opf_path = n
                break
    if not opf_path:
        for n in names:
            if n.endswith(".xhtml") or n.endswith(".html"):
                opf_path = n
                break
    if not opf_path:
        return ""
    try:
        oroot = ET.fromstring(z.read(opf_path))
    except Exception:
        return ""
    manifest = {}
    for it in oroot.iter():
        if _local(it.tag) == "item":
            mid, href = it.get("id"), it.get("href")
            if mid and href:
                manifest[mid] = href
    order = []
    for sp in oroot.iter():
        if _local(sp.tag) == "itemref":
            mid = sp.get("idref")
            if mid and mid in manifest:
                order.append(manifest[mid])
    if not order:
        order = list(manifest.values())
    base = opf_path.rsplit("/", 1)[0] if "/" in opf_path else ""
    out: List[str] = []
    for href in order:
        path = href if href.startswith("/") else (base + "/" + href if base else href)
        path = posixpath.normpath(path).lstrip("/")
        if path not in names:
            continue
        try:
            html = z.read(path).decode("utf-8", errors="ignore")
        except Exception:
            continue
        html = re.sub(r"<script[\s\S]*?</script>", " ", html, flags=re.I)
        html = re.sub(r"<style[\s\S]*?</style>", " ", html, flags=re.I)
        html = re.sub(r"<head[\s\S]*?</head>", " ", html, flags=re.I)
        html = re.sub(r"<title[\s\S]*?</title>", " ", html, flags=re.I)
        html = re.sub(r"<(p|div|br|/p|/div|/h[1-6]|li|/tr|/section|tr)[^>]*>", "\n", html, flags=re.I)
        text = re.sub(r"<[^>]+>", "", html)
        text = re.sub(r"&nbsp;", " ", text)
        text = re.sub(r"&amp;", "&", text)
        text = re.sub(r"&lt;", "<", text)
        text = re.sub(r"&gt;", ">", text)
        text = re.sub(r"\n{2,}", "\n", text).strip()
        if text:
            out.append(text)
    return "\n\n".join(out).strip()


def _extract_legacy_best_effort(raw: bytes) -> str:
    """对旧版 OLE(.doc/.xls) 做「尽力而为」文本提取（零依赖）。

    直接对二进制做 UTF-16LE / GBK 解码后抽取连续汉字/ASCII 游程，常能找回大部分正文
    （Word 正文多以 UTF-16LE 存于 WordDocument 流，或 ANSI 代码页以 GBK 存）。
    无有效文本返回空串；结果可能夹带少量格式乱码，属预期（系统提示已声明「尽力而为」）。
    """
    if not raw:
        return ""
    # 候选字符类：汉字/全角/半角字母数字/常用标点/空白，长度>=8
    pat = (r"[\u4e00-\u9fff\u3400-\u4dbf\u3000-\u303f\uff00-\uffef"
           r"A-Za-z0-9\s，。、；：！？（）「」“”‘’《》\-\._/]{8,}")
    cand: List[str] = []
    # 1) UTF-16LE（最常见）
    try:
        dec = raw.decode("utf-16le", errors="ignore")
        cand.extend(re.findall(pat, dec))
    except Exception:
        pass
    # 2) GBK / GB18030（旧版 ANSI 代码页中文）
    try:
        gbk = raw.decode("gb18030", errors="ignore")
        cand.extend(re.findall(pat, gbk))
    except Exception:
        pass
    # 3) 原始字节中的纯 ASCII 游程（标签/英文常以单字节存）
    for a in re.findall(rb"[\x20-\x7e]{8,}", raw):
        try:
            cand.append(a.decode("ascii"))
        except Exception:
            pass
    # 过滤：要求「含标点/数字/字母 且 中文占比>=50%」或长度>=24，
    # 剔除随机二进制解码出的伪中文游程（随机串中文占比低、且极少含数字/标点）
    def _keep(r: str) -> bool:
        if len(r) >= 24:
            return True
        cjk = sum(1 for c in r if "\u4e00" <= c <= "\u9fff" or "\u3400" <= c <= "\u4dbf")
        if cjk < len(r) * 0.5:
            return False
        return any(c in r for c in
                   "，。、；：！？（）0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz")
    ranked = [r.strip() for r in cand if len(r.strip()) >= 4 and _keep(r.strip())]
    ranked.sort(key=len, reverse=True)
    seen: set = set()
    keep: List[str] = []
    for r in ranked[:60]:
        key = r[:20]
        if key in seen:
            continue
        seen.add(key)
        keep.append(r)
    # 守卫：若无任何「实质连续文本」（最长游程 <12），视为无内容，回落报错
    if not keep or max(len(k) for k in keep) < 12:
        return ""
    return "\n".join(keep).strip()


def _extract_docx_library(raw: bytes) -> str:
    """第三方库回退（python-docx）。缺失依赖返回空串。"""
    try:
        from docx import Document  # type: ignore
    except Exception:
        return ""
    try:
        doc = Document(io.BytesIO(raw))
        lines: List[str] = []
        for para in doc.paragraphs:
            if para.text:
                lines.append(para.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text or "" for c in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))
        return "\n".join(lines).strip()
    except Exception:
        return ""


def _extract_xlsx_library(raw: bytes) -> str:
    """第三方库回退（openpyxl）。缺失依赖返回空串。"""
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception:
        return ""
    try:
        wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        lines: List[str] = []
        for ws in wb.worksheets:
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(cells):
                    rows.append("\t".join(cells))
            if rows:
                lines.append(f"### {ws.title}")
                lines.extend(rows)
        return "\n".join(lines).strip()
    except Exception:
        return ""


# 纯文本类扩展名（UTF-8 直接读取）
_TEXT_EXTS = {
    ".txt", ".text", ".md", ".markdown", ".csv", ".tsv", ".tab", ".json",
    ".jsonl", ".log", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
    ".xml", ".html", ".htm", ".py", ".js", ".ts", ".tsx", ".jsx", ".java",
    ".c", ".cpp", ".h", ".hpp", ".go", ".rs", ".rb", ".php", ".sh", ".bat",
    ".ps1", ".sql", ".r", ".scala", ".kt", ".swift", ".lua", ".pl", ".css",
    ".scss",
}

# 图片扩展名（进入对话作为视觉参考；系统不内置 OCR）
_IMAGE_EXTS_READ = {
    ".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tiff", ".svg",
}


def paginate_text(text: str, page: int = 1, page_size: int = 200,
                  max_chars: int = 8000, start_line: int = None,
                  end_line: int = None) -> dict:
    """把文本按行分页/切片，每行带行号（零填充对齐），并按字符预算截断。

    对标 hermes read_file 的「分页 / 行号 / 字符预算截断」能力，避免一次性把
    超大文档全文塞进上下文。

    两种模式（模型调用时参数名不统一，均兼容）：
      * 分页模式：page + page_size（默认第 1 页 200 行）。
      * 行范围模式：start_line + end_line（模型常用 read_file(..., start_line=200,
        end_line=400) 直接定位片段）；给定行范围时优先于分页。
    返回渲染所需的结构化信息。
    """
    lines = text.split("\n")
    total_lines = len(lines)
    width = max(3, len(str(total_lines)))

    # ── 行范围模式 ──
    if start_line is not None or end_line is not None:
        s = max(1, int(start_line or 1))
        e = int(end_line or total_lines)
        e = max(s, min(e, total_lines))
        s = min(s, e)
        slice_lines = lines[s - 1:e]
        numbered = "\n".join(
            f"{s + i:0{width}d}: {ln}" for i, ln in enumerate(slice_lines)
        ) or "(该范围为空)"
        truncated = False
        if len(numbered) > max_chars:
            numbered = numbered[:max_chars].rstrip() + "\n…(已按字符预算截断，调大 max_chars)"
            truncated = True
        return {
            "block": numbered,
            "total_lines": total_lines,
            "page": 1,
            "page_total": 1,
            "page_lines": len(slice_lines),
            "truncated": truncated,
            "range": (s, e),
        }

    # ── 分页模式 ──
    page = max(1, int(page or 1))
    page_size = max(1, int(page_size or 200))
    start = (page - 1) * page_size
    end = start + page_size
    page_lines = lines[start:end]
    numbered = "\n".join(
        f"{start + i + 1:0{width}d}: {ln}" for i, ln in enumerate(page_lines)
    ) or "(本页为空)"
    truncated = False
    if len(numbered) > max_chars:
        numbered = numbered[:max_chars].rstrip() + "\n…(已按字符预算截断，调大 max_chars 或翻页)"
        truncated = True
    page_total = max(1, (total_lines + page_size - 1) // page_size)
    return {
        "block": numbered,
        "total_lines": total_lines,
        "page": page,
        "page_total": page_total,
        "page_lines": len(page_lines),
        "truncated": truncated,
    }


def format_read(filename: str, ftype: str, pg: dict) -> str:
    if pg.get("range"):
        s, e = pg["range"]
        head = (
            f"文件 {filename} | 类型 {ftype} | 总行数 {pg['total_lines']} | "
            f"第 {s}-{e} 行 | 本段 {pg['page_lines']} 行"
        )
    else:
        head = (
            f"文件 {filename} | 类型 {ftype} | 总行数 {pg['total_lines']} | "
            f"第 {pg['page']}/{pg['page_total']} 页 | 本页 {pg['page_lines']} 行"
        )
    tail = ""
    if pg["truncated"] or pg.get("range") or pg["page"] < pg["page_total"]:
        tail = "\n（内容较长：用 start_line/end_line 或 page 参数定位，max_chars 控制返回长度）"
    return head + "\n" + "-" * 40 + "\n" + pg["block"] + tail


def _split_text(text: str, chunk_size: int = 400, overlap: int = 50) -> List[str]:
    """按中英文标点切分，滑动窗口重叠。"""
    paras = [p.strip() for p in re.split(r"\n+", text) if p.strip()]
    chunks, buf = [], ""
    for p in paras:
        if len(buf) + len(p) <= chunk_size:
            buf += p + "\n"
        else:
            if buf:
                chunks.append(buf.strip())
            if len(p) > chunk_size:
                for i in range(0, len(p), chunk_size - overlap):
                    chunks.append(p[i:i + chunk_size].strip())
                buf = ""
            else:
                buf = p + "\n"
    if buf:
        chunks.append(buf.strip())
    return [c for c in chunks if c]


def _ext(path: str) -> str:
    return os.path.splitext(path)[1].lower()


def _is_text_ext(name: str) -> bool:
    """文件名（含路径/扩展名）是否属文本类。"""
    return os.path.splitext(name)[1].lower() in _TEXT_EXTS


def _looks_like_text(raw: bytes, threshold: float = 0.7) -> bool:
    """字节流中可打印字符(含常见空白)占比是否足够高，用于判断是否为可读文本。

    避免把『高可读性字节流』（如 .dat/.bin 实为文本）误判为二进制而报错。
    """
    if not raw:
        return False
    sample = raw[:8192]
    printable = sum(1 for b in sample if 32 <= b < 127 or b in (9, 10, 13))
    return (printable / len(sample)) >= threshold


def _looks_like_zip(raw: bytes) -> bool:
    """Magic bytes 检测 ZIP 容器：本地文件头 / 中央目录尾 / 空归档 / 分卷 (PK...)。"""
    return raw[:4] in (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")


# 以下扩展名内部虽是 ZIP 容器，但属于有专用解析器的办公文档，
# 绝不能走通用压缩包解包分支（否则会吐出原始 XML 标签噪声）。
_ZIP_BASED_DOC_EXTS = (".docx", ".xlsx", ".pptx", ".odt", ".ods", ".odp", ".epub")


def _sniff_office_ext(raw: bytes) -> str | None:
    """通过 magic bytes + 容器结构嗅探办公文档真实类型，纠错错配/缺失的扩展名。

    返回 '.docx' / '.xlsx' / '.pptx' / '.odt' / '.ods' / '.odp' / '.epub'
    / '.doc' / '.xls' 之一，或 None（不是可识别的办公文档）。只读探测，不解析正文。
    """
    if not raw:
        return None
    # 旧版 OLE 二进制（.doc / .xls / .ppt），非 ZIP
    if raw[:4] == b"\xd0\xcf\x11\xe0":
        if b"PowerPoint Document" in raw or b"PowerPoint" in raw:
            return ".ppt"
        if b"WordDocument" in raw:
            return ".doc"
        if b"Workbook" in raw or b"Book" in raw:
            return ".xls"
        return ".doc"  # 兜底：走 _extract_legacy_best_effort 尽力而为
    if not _looks_like_zip(raw):
        return None
    try:
        z = zipfile.ZipFile(io.BytesIO(raw))
        names = set(z.namelist())
    except Exception:
        return None
    # OpenDocument：mimetype 文件标识
    try:
        if "mimetype" in names:
            mt = z.read("mimetype").decode("utf-8", "ignore").strip()
            if mt == "application/vnd.oasis.opendocument.text":
                return ".odt"
            if mt == "application/vnd.oasis.opendocument.spreadsheet":
                return ".ods"
            if mt == "application/vnd.oasis.opendocument.presentation":
                return ".odp"
        # EPUB：container.xml + OEBPS 内容目录
        if "META-INF/container.xml" in names and any(
                n.startswith("OEBPS/") for n in names):
            return ".epub"
    except Exception:
        pass
    # OOXML
    if "word/document.xml" in names:
        return ".docx"
    joined = " ".join(names)
    if "xl/workbook.xml" in names or "xl/worksheets" in joined:
        return ".xlsx"
    if "ppt/presentation.xml" in names or "ppt/slides" in joined:
        return ".pptx"
    return None


def _extract_zip(raw: bytes, depth: int = 0) -> str:
    """从 ZIP 压缩包递归提取可读文本（零依赖标准库，对标 Hermes 自愈式解析）。

    这是『按需读取 + 自愈』的核心：遇到压缩包不报错甩锅，而是自动解包、
    列出内部条目、递归提取文本类文件内容，让 Agent 直接拿到可分析的文本。
    若无可提取文本则仅返回清单，交由上层决定如何提示。
    """
    try:
        z = zipfile.ZipFile(io.BytesIO(raw))
    except Exception:
        return ""
    entries = z.namelist()
    if not entries:
        return ""
    out = [f"[压缩包内含 {len(entries)} 个条目]\n"] if depth == 0 else []
    for name in entries:
        if name.endswith("/"):
            continue  # 目录条目
        try:
            data = z.read(name)
        except Exception:
            continue
        # 嵌套压缩包：递归解包（最多 2 层，防止极深结构耗尽资源）
        if _looks_like_zip(data) and depth < 2:
            nested = _extract_zip(data, depth + 1)
            if nested.strip():
                out.append(f"\n=== 嵌套压缩包 {name} ===\n{nested}")
            continue
        if _is_text_ext(name) or _looks_like_text(data):
            try:
                content = data.decode("utf-8", errors="ignore")
            except Exception:
                content = ""
            if content.strip():
                out.append(f"\n--- {name} ---\n{content}")
    return "\n".join(out)


def _extract_pdf(raw: bytes) -> str:
    """从 PDF 二进制提取文本（跨平台纯 Python 方案，不依赖任何 OCR 引擎）。

    策略：
      1) PyMuPDF(fitz)：逐页提取文本层（文本型 PDF 的标准做法，跨平台一致）；
      2) pypdf：若已安装则作为兜底（单页容错 + 空口令解密）；
      3) pdfminer.six：终极兜底，对中文/CID 字体更鲁棒；
      4) 三者都失败（通常是纯图片扫描件、无文字层）返回空字符串，
         交由 read_file_text 渲染为图片或提示用户——本系统刻意不内置 OCR
         （tesseract / PaddleOCR 等需原生二进制，跨平台部署不可控）。
    """
    # 1) PyMuPDF（环境已安装）为主提取器 —— 仅取文本层，不做 OCR
    try:
        import fitz  # type: ignore
        doc = fitz.open(stream=io.BytesIO(raw), filetype="pdf")
        parts: list = []
        for pg in doc:
            try:
                t = pg.get_text("text") or ""
            except Exception:
                t = ""
            if t.strip():
                parts.append(t)
        doc.close()
        text = "\n\n".join(parts).strip()
        if text:
            return text
    except Exception:
        pass

    # 2) pypdf 兜底
    try:
        from pypdf import PdfReader  # type: ignore
        reader = PdfReader(io.BytesIO(raw))
        try:
            reader.decrypt("")  # 兼容空用户口令加密的文档
        except Exception:
            pass
        parts = []
        for pg in reader.pages:
            try:
                parts.append(pg.extract_text() or "")
            except Exception:
                parts.append("")
        text = "\n".join(parts).strip()
        if text:
            return text
    except Exception:
        pass

    # 3) pdfminer 兜底
    try:
        from pdfminer.high_level import extract_text  # type: ignore
        text = extract_text(io.BytesIO(raw)).strip()
        if text:
            return text
    except Exception:
        pass

    return ""


# ── 图片 / 扫描件 OCR（Tesseract，需系统安装 tesseract-ocr + tesseract-ocr-chi-sim）──
# 缺失引擎时函数返回空串，调用方降级（图片→提示视觉参考；扫描 PDF→渲染为图）。
_IMG_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tiff", ".tif", ".svg"}

# LibreOffice 无头转换全局串行锁：soffice 不允许并发实例（UserInstallation 配置目录锁），
# 多上传并发时必须串行化，否则后到者会直接失败。
_LO_LOCK = threading.Lock()
# 单次 OCR / 转换的页数上限（防超大扫描件长时间阻塞事件循环）
_OCR_MAX_PAGES = 30


def _libreoffice_convert(raw: bytes, filename: str, target_ext: str,
                         media_dir=None, owner=None) -> Tuple[str, str] | None:
    """用 LibreOffice 无头模式把旧版 OLE 文档(.doc/.xls/.ppt) 转成新格式再解析。

    返回 (text, file_type) 或 None（无 soffice / 转换失败）。转换成功后递归调用
    read_file_text 走既有标准库/第三方提取器，零额外解析逻辑、不重复造轮子。
    全局 threading.Lock 串行化，避免 soffice 实例锁冲突。
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    tmp = tempfile.mkdtemp(prefix="zhishu_lo_")
    try:
        in_path = os.path.join(tmp, os.path.basename(filename) or "input")
        with open(in_path, "wb") as f:
            f.write(raw)
        out_dir = os.path.join(tmp, "out")
        os.makedirs(out_dir, exist_ok=True)
        prof = os.path.join(tmp, "profile")
        os.makedirs(prof, exist_ok=True)
        cmd = [
            soffice, "--headless", "--norestore", "--nofirststartwizard",
            f"-env:UserInstallation=file://{prof}",
            "--convert-to", target_ext.lstrip("."), "--outdir", out_dir, in_path,
        ]
        with _LO_LOCK:
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
        if proc.returncode != 0:
            return None
        out_files = [os.path.join(out_dir, f) for f in os.listdir(out_dir)
                     if f.lower().endswith(target_ext.lower())]
        if not out_files:
            return None
        with open(out_files[0], "rb") as f:
            out_raw = f.read()
        # 递归解析转换后的新格式（.docx/.xlsx/.pptx），不触发二次转换
        return read_file_text(os.path.basename(out_files[0]), out_raw, media_dir, owner)
    except Exception:
        return None
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _ocr_image_bytes(raw: bytes, lang: str = "chi_sim+eng") -> str:
    """对图片二进制做 OCR（需 tesseract + 中文包）。失败/未安装返回空串。"""
    try:
        from PIL import Image
        import pytesseract
    except Exception:
        return ""
    try:
        img = Image.open(io.BytesIO(raw))
        img = img.convert("RGB")
        return pytesseract.image_to_string(img, lang=lang).strip()
    except Exception:
        return ""


def _ocr_pdf(raw: bytes, lang: str = "chi_sim+eng", max_pages: int = _OCR_MAX_PAGES) -> str:
    """对扫描型 PDF（无文字层）逐页渲染后 OCR。失败/未安装返回空串。"""
    try:
        import fitz  # type: ignore
        from PIL import Image
        import pytesseract
    except Exception:
        return ""
    try:
        doc = fitz.open(stream=io.BytesIO(raw), filetype="pdf")
        parts: list = []
        for i, pg in enumerate(doc):
            if i >= max_pages:
                break
            pix = pg.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            t = pytesseract.image_to_string(img, lang=lang).strip()
            if t:
                parts.append(t)
        doc.close()
        return "\n\n".join(parts).strip()
    except Exception:
        return ""


def read_file_text(filename: str, raw: bytes,
                   media_dir: str | None = None,
                   owner: str | None = None) -> tuple[str, str]:
    """从上传的二进制内容提取纯文本。

    返回 (text, file_type)。file_type 用于前端展示标签。
    无法解析时抛出 ValueError（带友好中文说明）。
    """
    ext = _ext(filename)
    # 纠错：扩展名缺失/错配但 magic bytes 表明是办公文档时，按真实类型派发，
    # 避免被下方通用压缩包分支当成 ZIP 吐出原始 XML。
    if ext not in _ZIP_BASED_DOC_EXTS and ext != ".zip":
        sniffed = _sniff_office_ext(raw)
        if sniffed:
            ext = sniffed
    file_type = ext.lstrip(".").upper() or "TXT"

    # ── 压缩包：自动解包并递归提取内部文本（自愈核心，而非报错甩锅）──
    # 即使文件被错配扩展名（如 .txt 实为 zip），也按 magic bytes 识别处理。
    # 但 .docx/.xlsx/.pptx/.odt/.ods/.odp/.epub 内部也是 ZIP，需先走专用解析器，
    # 故当且仅当扩展名不在办公文档集合时才进入通用解包分支。
    if (ext == ".zip" or _looks_like_zip(raw)) and ext not in _ZIP_BASED_DOC_EXTS:
        text = _extract_zip(raw)
        if text.strip():
            return text, "ZIP"
        # 解包无可读文本：仍返回清单 + 提示（让 Agent 知道这是归档而非二进制）
        return ("[压缩包] " + text + "\n"
                "(归档内无可直接读取的文本文件；如需处理内部二进制，请用 code_exec 解包："
                "import zipfile,os; z=zipfile.ZipFile(os.environ['TARGET_FILE']); print(z.namelist())"), "ZIP"

    if ext in _TEXT_EXTS:
        return raw.decode("utf-8", errors="ignore"), file_type or "TXT"

    if ext == ".pdf":
        text = _extract_pdf(raw)
        if not text.strip():
            # 文本层为空：先尝试 OCR（扫描件/图片型 PDF，需系统 tesseract + 中文包）
            ocr = _ocr_pdf(raw)
            if ocr.strip():
                return ocr, "PDF(OCR)"
            # OCR 也失败：扫描件渲染为图片，至少可见内容，避免「无结果」死路。
            try:
                import fitz  # type: ignore
                doc = fitz.open(stream=io.BytesIO(raw), filetype="pdf")
                # 渲染图按用户命名空间隔离：media_dir/<owner>/pdf_render/...，
                # 经由 main.py 的 /media 网关按 owner 段位校验，杜绝跨用户读取他人扫描件。
                if media_dir and owner:
                    base = os.path.join(media_dir, owner, "pdf_render")
                    url_prefix = "/media/" + owner + "/pdf_render/"
                elif media_dir:
                    base = os.path.join(media_dir, "_shared", "pdf_render")
                    url_prefix = "/media/_shared/pdf_render/"
                else:
                    base = os.path.join("data", "generated", "pdf_render")
                    url_prefix = "/media/pdf_render/"
                os.makedirs(base, exist_ok=True)
                safe = re.sub(r"[^\w\-.]", "_", os.path.basename(filename)) or "doc"
                out_dir = os.path.join(base, safe + "_" + uuid.uuid4().hex[:8])
                os.makedirs(out_dir, exist_ok=True)
                media_urls = []
                for i, pg in enumerate(doc):
                    pix = pg.get_pixmap(dpi=150)
                    p = os.path.join(out_dir, f"page_{i + 1}.png")
                    pix.save(p)
                    media_urls.append(
                        url_prefix + os.path.basename(out_dir) + f"/page_{i + 1}.png")
                doc.close()
                if media_urls:
                    raise ValueError(
                        "该 PDF 为纯图片扫描件（无文字层）。系统为保持跨平台一致、不内置 OCR 识别，"
                        "无法直接提取文字。已为您将各页渲染为图片，可在前端/以下地址查看内容：\n"
                        + "\n".join(media_urls)
                        + "\n（如需文字版，请提供由文字处理软件导出的文本型 PDF，或先转换为可提取文本的格式。）")
            except ValueError:
                raise
            except Exception:
                pass
            raise ValueError(
                "PDF 未提取到文本：文件可能已加密、受损，或为纯图片扫描件（无文字层）。"
                "系统不内置 OCR 识别（跨平台兼容性考虑）；若为扫描件请提供文本型 PDF。")
        return text, "PDF"

    if ext in (".docx", ".doc"):
        if ext == ".doc":
            # 旧版 .doc 为 OLE 二进制：优先 LibreOffice 无头转 docx（保真高），
            # 失败再尽力而为字节扫描，仍失败才引导转存。
            conv = _libreoffice_convert(raw, filename, ".docx", media_dir, owner)
            if conv and conv[0].strip():
                return conv[0], "DOC"
            text = _extract_legacy_best_effort(raw)
            if text.strip():
                return text, "DOC"
            raise ValueError(
                "暂不支持旧版 .doc 格式（OLE 二进制结构，标准库无法完整解析）。"
                "已尝试 LibreOffice 转换与字节扫描，均未获得有效文本。"
                "可选方案：① 请用户另存为 .docx 后重新上传以获得完整解析；"
                "② 或调用 code_exec 对文件做「尽力而为」的文本提取（扫描 UTF-16LE / 可打印字符，"
                "可能仅恢复部分正文），路径通过 os.environ['TARGET_FILE'] 读取。"
            )
        # 优先零依赖标准库解析；为空再回退第三方库；仍为空则视为图片型文档
        text = _extract_docx_stdlib(raw) or _extract_docx_library(raw)
        if not text.strip():
            raise ValueError("Word 文档未提取到文本（可能是图片型文档；可尝试 OCR，但当前未提取到文字，请转换为文本型文档）。")
        return text, "DOCX"

    if ext in (".xlsx",):
        # 优先零依赖标准库解析；为空再回退第三方库
        text = _extract_xlsx_stdlib(raw) or _extract_xlsx_library(raw)
        if not text.strip():
            raise ValueError("Excel 文档未提取到内容")
        return text, "XLSX"

    if ext == ".xls":
        # 旧版 .xls 为 OLE 二进制：优先 LibreOffice 无头转 xlsx，失败再字节扫描
        conv = _libreoffice_convert(raw, filename, ".xlsx", media_dir, owner)
        if conv and conv[0].strip():
            return conv[0], "XLS"
        text = _extract_legacy_best_effort(raw)
        if text.strip():
            return text, "XLS"
        raise ValueError(
            "暂不支持旧版 .xls 格式（OLE 二进制结构，非 zip，标准库无法解析）。"
            "已尝试 LibreOffice 转换与字节扫描，均未获得有效文本。"
            "可选方案：① 请用户另存为 .xlsx 后重新上传以获得完整解析；"
            "② 或调用 code_exec 对文件做「尽力而为」的文本提取（扫描可打印字符），"
            "路径通过 os.environ['TARGET_FILE'] 读取。"
        )

    if ext == ".ppt":
        # 旧版 .ppt 为 OLE 二进制：LibreOffice 无头转 pptx 后按新格式解析
        conv = _libreoffice_convert(raw, filename, ".pptx", media_dir, owner)
        if conv and conv[0].strip():
            return conv[0], "PPT"
        raise ValueError(
            "暂不支持旧版 .ppt 格式（OLE 二进制结构，非 zip，标准库无法解析）。"
            "已尝试 LibreOffice 转换未成功。可选方案：请用户另存为 .pptx 后重新上传以获得完整解析。"
        )

    if ext == ".pptx":
        text = _extract_pptx_stdlib(raw)
        if not text.strip():
            raise ValueError("PPT 文档未提取到文本（可能是图片型幻灯片）")
        return text, "PPTX"

    if ext in (".odt", ".ods", ".odp"):
        text = _extract_odf_stdlib(raw)
        if not text.strip():
            raise ValueError("OpenDocument 文档未提取到文本")
        return text, ext.lstrip(".").upper()

    if ext == ".rtf":
        text = _extract_rtf(raw)
        if not text.strip():
            raise ValueError("RTF 文档未提取到文本")
        return text, "RTF"

    if ext == ".epub":
        text = _extract_epub_stdlib(raw)
        if not text.strip():
            raise ValueError("EPUB 未提取到文本")
        return text, "EPUB"

    if ext in _IMG_EXTS:
        # 图片：OCR 提取文字（需系统 tesseract + 中文包）；无文字则提示作为视觉参考
        ocr = _ocr_image_bytes(raw)
        if ocr.strip():
            return ocr, "IMAGE(OCR)"
        raise ValueError(
            f"图片「{filename}」未识别到文字（可能是无文字的图表/照片）。"
            "已尝试 OCR 但未提取到文本；请作为视觉参考(vision)传入模型以获得内容理解。"
        )

    # 兜底：高可读性字节流直接当文本；GBK/GB18030 兜底；否则给出可操作的 code_exec 自愈指令
    if _looks_like_text(raw):
        return raw.decode("utf-8", errors="ignore"), file_type or "TXT"
    try:
        gbk = raw.decode("gb18030")
        if gbk.strip():
            return gbk, file_type or "TXT"
    except Exception:
        pass
    # 死路保护：把『报错甩锅』变成『可执行的自愈指引』，引导 Agent 用 code_exec 自救
    raise ValueError(
        f"文件「{filename}」（扩展名 {ext or '未知'}）无法直接解析为文本。\n"
        "请勿直接告诉用户『无法解析』或让其重新上传；应先调用 code_exec 编写 Python 自救：\n"
        "  import os\n"
        "  p = os.environ.get('TARGET_FILE')  # stored_path 对应的磁盘文件\n"
        "  data = open(p, 'rb').read()\n"
        "  # 据字节特征（如 PK 头=ZIP、\\x1f\\x8b=Gzip、7z\\xbc\\xaf=7z）选择对应库解析\n"
        "  print(len(data), data[:16].hex())  # 先探查 magic bytes，再决定处理方式")


class KnowledgeBase:
    def __init__(self, emb_cfg: EmbeddingConfig, vs_cfg: VectorStoreConfig,
                 data_dir: Optional[str] = None, app_cfg=None):
        self.emb = EmbeddingEngine(emb_cfg, app_cfg=app_cfg)
        # 向量库路径必须与其它持久化（用户/审计/记忆/会话库）一样跟随 data_dir。
        # 旧实现直接用 vs_cfg.path 这个**相对当前工作目录**的裸路径，导致：
        #   * 换个 CWD 启动 → 其它库仍在 data_dir，向量库却跑到新目录并被建成空库，
        #     表现为「知识库突然全空」，排查成本极高；
        #   * 测试/多实例无法通过 data_dir 隔离，会串到真实运行数据上。
        # 兼容性：默认配置 data_dir=data、path=data/zhishu_vector.db，去掉重复的
        # 前导 data/ 后再拼接，结果仍是 data/zhishu_vector.db —— 存量部署零变化。
        if data_dir:
            vs_cfg.path = self._resolve_store_path(vs_cfg.path, data_dir)
        self.store = VectorStore(vs_cfg)
        # 原始文件保留目录（用于「重新解析」）。仅当传入 data_dir 时启用。
        self.raw_dir = os.path.join(data_dir, "knowledge_raw") if data_dir else ""
        if self.raw_dir:
            os.makedirs(self.raw_dir, exist_ok=True)
        # 知识图谱层（关键词共现网络，离线）。data_dir 解析为绝对路径，
        # 与向量库平行持久化于 data_dir/zhishu_kg.db。
        self.data_dir = os.path.abspath(data_dir) if data_dir else None
        # 媒体托管目录：扫描件 PDF 渲染图写入此处，经 /media 同源托管
        self.media_dir = None
        if self.data_dir and app_cfg is not None:
            _store = getattr(getattr(app_cfg, "media", None), "store_dir", None)
            if _store:
                self.media_dir = os.path.join(self.data_dir, _store)
        self.graph = KnowledgeGraph(self.data_dir)

    @staticmethod
    def _resolve_store_path(path: str, data_dir: str) -> str:
        """把向量库路径归一到 data_dir 下（绝对路径原样保留）。"""
        p = (path or "").strip() or "zhishu_vector.db"
        # 注意：Windows 的 os.path.isabs("/data/kb.db") 为 False（无盘符不算绝对），
        # 但配置里写 POSIX 绝对路径的意图很明确，必须原样保留，否则会被拼到
        # data_dir 下变成 data/data/kb.db 之类的错误位置。
        if os.path.isabs(p) or p.startswith(("/", "\\")):
            return p
        parts = p.replace("\\", "/").split("/")
        # 去掉历史默认前缀 "./data/" / "data/"，避免拼出 data/data/xxx.db
        while parts and parts[0] in (".", "data"):
            parts.pop(0)
        if not parts:
            parts = ["zhishu_vector.db"]
        return os.path.join(data_dir, *parts)

    # ------------------------- 入库（带归属与元数据） -------------------------
    def ingest_text(
        self,
        text: str,
        doc_id: str = None,
        meta: dict = None,
        owner: str = None,
        title: str = None,
        file_type: str = "TEXT",
        source: str = None,
    ) -> dict:
        doc_id = doc_id or uuid.uuid4().hex[:12]
        chunks = _split_text(text)
        if not chunks:
            return {"doc_id": doc_id, "chunks": 0, "skipped": True}
        # 降级隔离：向量随「实际使用的后端」打签名。Provider 抖动时产生的 hash
        # 伪向量与真语义向量不在同一空间，混存会污染检索（维度不同还会抛异常）。
        # strict_ingest=True 时宁可入库失败也不写脏数据。
        vecs, emb_sig, degraded = self.emb.embed_tagged(chunks)
        if degraded and getattr(self.emb.cfg, "strict_ingest", True):
            raise ValueError(
                "入库已中止：当前 embedding 模型不可用，系统本会降级为 hash 伪向量，"
                "而伪向量与既有语义向量不在同一空间，写入会污染知识库检索。"
                "请在「模型管理」中确认 embedding 模型（embedding.embed_model）可用后重试；"
                "若确需先入库，可将 embedding.strict_ingest 设为 false（该文档需在模型"
                "恢复后重新解析才能被语义检索到）。"
            )
        full_meta = {
            "title": title or doc_id,
            "source": source or title or doc_id,
            "file_type": file_type,
            "owner": owner,
            "size": len(text.encode("utf-8")),
            "char_count": sum(len(c) for c in chunks),
            "content_preview": text,
            **(meta or {}),
        }
        n = self.store.add(doc_id, chunks, vecs, full_meta, emb_sig=emb_sig)
        # 增量构建知识图谱（关键词共现网络）
        if self.graph is not None:
            try:
                self.graph.analyze_document(doc_id, text, owner=owner)
            except Exception as e:  # 图谱失败绝不影响主链路
                import logging
                logging.getLogger("zhishu.rag").warning("知识图谱分析失败 doc_id=%s: %s", doc_id, e)
        out = {"doc_id": doc_id, "chunks": n, "title": title or doc_id,
               "file_type": file_type}
        if degraded:
            out["degraded"] = True
            out["warning"] = (
                "embedding 模型不可用，本次已使用 hash 伪向量入库：该文档在语义"
                "检索中不可见，请在模型恢复后对其执行「重新解析」。")
        return out

    def ingest_file(
        self,
        filename: str,
        raw: bytes,
        doc_id: str = None,
        owner: str = None,
        title: str = None,
    ) -> dict:
        """从上传的二进制内容解析并入库。失败时抛出 ValueError。"""
        text, file_type = read_file_text(filename, raw, self.media_dir, owner)
        if not doc_id:
            # 安全：doc_id 以「owner+文件名」哈希命名空间化 —— 不同用户上传同名
            # 文件不会碰撞（旧实现直接用文件名，B 上传同名文件会顶掉 A 的文档
            # 归属并造成向量串库）；同一用户重复上传同名文件仍覆盖自己的旧文档。
            import hashlib
            stem = os.path.splitext(filename)[0] or filename or "doc"
            doc_id = hashlib.sha1(
                f"{owner or ''}:{stem}".encode("utf-8")).hexdigest()[:16]
        raw_path = self._save_raw(doc_id, filename, raw)
        return self.ingest_text(
            text,
            doc_id=doc_id,
            meta={"raw_path": raw_path} if raw_path else {},
            owner=owner,
            title=title or os.path.splitext(filename)[0] or filename,
            file_type=file_type,
            source=filename,
        )

    def _save_raw(self, doc_id: str, filename: str, raw: bytes) -> str:
        """将原始字节落盘，供后续「重新解析」使用。失败返回空串。"""
        if not self.raw_dir:
            return ""
        safe = re.sub(r"[^\w\-.]", "_", os.path.basename(filename))
        path = os.path.join(self.raw_dir, f"{doc_id}__{safe}")
        try:
            with open(path, "wb") as f:
                f.write(raw)
            return os.path.abspath(path)
        except OSError:
            return ""

    def reparse_document(self, doc_id: str, owner: Optional[str] = None) -> dict:
        """用保留的原始文件，以当前解析器重新提取并覆盖入库（保持 doc_id 不变）。"""
        if not self.raw_dir:
            raise ValueError("服务端未启用原始文件保留，无法重新解析")
        doc = self.store.get_document(doc_id, owner=owner)
        if not doc:
            raise ValueError("文档不存在或无权限")
        raw_path = doc.get("raw_path")
        if not raw_path or not os.path.exists(raw_path):
            raise ValueError(
                "未保留原始文件，无法重新解析（可能是早期文本入库或旧版上传）")
        with open(raw_path, "rb") as f:
            raw = f.read()
        filename = doc.get("source") or f"{doc_id}.bin"
        try:
            text, file_type = read_file_text(filename, raw, self.media_dir, owner)
        except ValueError as e:
            raise ValueError(f"重新解析失败：{e}")
        if not text.strip():
            raise ValueError("重新解析后内容为空（可能是图片型文档；系统不内置 OCR，无法提取图片内文字，请转换为文本型文档）。")
        # 删除旧分块与元数据（delete_document 会一并清理 raw_path 文件）
        self.store.delete_document(doc_id, owner=owner)
        # 回退旧图谱贡献（reparse 走的是 store.delete，不会触发 KB.delete_document）
        if self.graph is not None:
            try:
                self.graph.remove_document(doc_id)
            except Exception:
                pass
        # 重新写回原始文件，保证可再次重新解析
        try:
            with open(raw_path, "wb") as f:
                f.write(raw)
        except OSError:
            raw_path = ""
        res = self.ingest_text(
            text,
            doc_id=doc_id,
            meta={"raw_path": raw_path} if raw_path else {},
            owner=owner,
            title=doc.get("title"),
            file_type=file_type,
            source=doc.get("source"),
        )
        return {
            "doc_id": doc_id,
            "title": doc.get("title"),
            "file_type": file_type,
            "chunks": res.get("chunks", 0),
        }

    def ingest_local_file(self, path: str, doc_id: str = None,
                          owner: str = None) -> dict:
        with open(path, "rb") as f:
            raw = f.read()
        return self.ingest_file(os.path.basename(path), raw, doc_id, owner)

    # ------------------------- 检索 / 上下文 -------------------------
    def query(self, question: str, top_k: int = 5,
              owner: Optional[str] = None) -> List[dict]:
        # 检索侧同样按签名隔离：用 hash 伪向量去比对真语义向量（或反之）只会
        # 得到噪声排序，且维度不同会抛异常。签名不匹配的分块直接不参与打分。
        vecs, emb_sig, degraded = self.emb.embed_tagged([question])
        if degraded:
            import logging
            logging.getLogger("zhishu.rag").warning(
                "检索时 embedding 已降级为 hash，仅能命中同为 hash 签名的分块；"
                "语义检索结果将显著变差，请检查 embedding 模型可用性。")
        return self.store.search(vecs[0], top_k, owner=owner, emb_sig=emb_sig)

    def build_context(self, question: str, top_k: int = 5,
                      owner: Optional[str] = None) -> str:
        hits = self.query(question, top_k, owner=owner)
        if not hits:
            return ""
        parts = []
        for i, h in enumerate(hits, 1):
            src = h["meta"].get("source", h["doc_id"])
            parts.append(f"[知识 {i} | 来源:{src} | 相似度:{h['score']:.3f}]\n{h['text']}")
        return "\n\n".join(parts)

    # ------------------------- 文档级管理 -------------------------
    def list_documents(self, owner: Optional[str] = None,
                       limit: int = 200, offset: int = 0,
                       q: Optional[str] = None) -> List[dict]:
        return self.store.list_documents(owner, limit, offset, q)

    def get_document(self, doc_id: str, owner: Optional[str] = None) -> Optional[dict]:
        return self.store.get_document(doc_id, owner)

    def delete_document(self, doc_id: str, owner: Optional[str] = None) -> bool:
        ok = self.store.delete_document(doc_id, owner)
        if ok and self.graph is not None:
            try:
                self.graph.remove_document(doc_id)
            except Exception:
                pass
        return ok

    def stats(self, owner: Optional[str] = None) -> dict:
        out = {
            "backend": self.store.backend,
            "embedding_dim": self.emb.dim,
            "vectors": self.store.count(owner),
            "documents": self.store.doc_count(owner),
        }
        # 暴露向量空间签名与「陈旧分块」数量：换了 embedding 模型或曾经降级过的
        # 分块在当前配置下检索不到，前端可据此提示用户重新解析。
        try:
            sig = self.emb.signature
            out["embedding_signature"] = sig
            out.update(self.store.signature_stats(sig))
        except Exception:  # noqa: BLE001 —— 统计失败不应影响主接口
            pass
        return out
