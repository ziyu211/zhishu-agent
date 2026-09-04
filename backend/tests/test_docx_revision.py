"""docx_revision 修订版生成器回归测试。

背景：生成「带 Word 原生批注的修订版 docx」此前完全交给模型手写 oxml，
实测 240 份产物成功率仅 22%（188 份无 comments.xml，其中 3 份是 22 字节空 zip）。
本测试锁定服务端实现的正确性：批注三件套齐全 + zip 健康 + 修订样式正确。
"""
from __future__ import annotations

import os
import sys
import tempfile
import zipfile
import xml.etree.ElementTree as ET

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from zhishu.core.docx_revision import (  # noqa: E402
    build_revised_docx, DocxRevisionError, _REL_COMMENTS,
)

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _make_docx(path, paragraphs, bold=False):
    """生成一个测试用 docx。"""
    from docx import Document
    from docx.shared import Pt
    d = Document()
    for text in paragraphs:
        p = d.add_paragraph()
        r = p.add_run(text)
        if bold:
            r.bold = True
        r.font.size = Pt(12)
    d.save(path)
    return path


def _comments(zf):
    root = ET.fromstring(zf.read("word/comments.xml"))
    return [e for e in root if e.tag == W + "comment"]


def _comment_texts(zf):
    out = []
    for c in _comments(zf):
        out.append("".join(t.text or "" for t in c.iter(W + "t")))
    return out


# ─────────────────────────── 核心用例 ───────────────────────────

def test_basic_revision_with_comment():
    """基础：错字红色删除线 + 正字红色 + 一条原生批注。"""
    with tempfile.TemporaryDirectory() as td:
        src = _make_docx(os.path.join(td, "in.docx"), ["提升居民消废体验"])
        data, st = build_revised_docx(src, [{"find": "消废", "replace": "消费"}])

        out = os.path.join(td, "out.docx")
        open(out, "wb").write(data)
        zf = zipfile.ZipFile(out)

        # 1) zip 健康
        assert zf.testzip() is None
        # 2) 批注部件存在且条数正确
        assert "word/comments.xml" in zf.namelist()
        assert len(_comments(zf)) == 1
        assert "消废" in _comment_texts(zf)[0] and "消费" in _comment_texts(zf)[0]
        # 3) 三件套：锚点 + Content_Types + rels
        dx = zf.read("word/document.xml").decode("utf-8")
        assert "commentRangeStart" in dx and "commentReference" in dx
        assert "/word/comments.xml" in zf.read("[Content_Types].xml").decode("utf-8")
        assert _REL_COMMENTS in zf.read("word/_rels/document.xml.rels").decode("utf-8")
        assert st["applied"] == 1 and st["comments"] == 1 and not st["missed"]


def test_revision_styles_strike_and_red():
    """错字 run 必须同时具备红色与删除线；正字 run 仅红色。"""
    with tempfile.TemporaryDirectory() as td:
        src = _make_docx(os.path.join(td, "in.docx"), ["提升居民消废体验"])
        data, _ = build_revised_docx(src, [{"find": "消废", "replace": "消费"}])
        out = os.path.join(td, "o.docx")
        open(out, "wb").write(data)
        zf = zipfile.ZipFile(out)
        root = ET.fromstring(zf.read("word/document.xml"))

        runs = root.iter(W + "r")
        strike_red, plain_red = [], []
        for r in runs:
            txt = "".join(t.text or "" for t in r.iter(W + "t"))
            if txt not in ("消废", "消费"):
                continue
            rPr = r.find(W + "rPr")
            has_strike = rPr is not None and rPr.find(W + "strike") is not None
            color = ""
            if rPr is not None:
                c = rPr.find(W + "color")
                if c is not None:
                    color = (c.get(W + "val") or "").upper()
            (strike_red if has_strike else plain_red).append((txt, color))

        assert any(t == "消废" and c == "FF0000" for t, c in strike_red), strike_red
        assert any(t == "消费" and c == "FF0000" for t, c in plain_red), plain_red
        # 正字不应带删除线
        assert all(t != "消费" for t, _c in strike_red)


def test_delete_extra_char():
    """多字（删除多余字）：replace 为空串 → 只留红色删除线，不补正字。"""
    with tempfile.TemporaryDirectory() as td:
        src = _make_docx(os.path.join(td, "in.docx"),
                         ["集成 DID 分布式数字身份、联盟链、隐私计算机等技术"])
        data, st = build_revised_docx(
            src, [{"find": "计算机", "replace": "计算",
                   "comment": "删除多余「机」字：与并列项同构应为技术名词「隐私计算」"}])
        out = os.path.join(td, "o.docx")
        open(out, "wb").write(data)
        zf = zipfile.ZipFile(out)
        assert zf.testzip() is None
        assert len(_comments(zf)) == 1
        assert "删除多余" in _comment_texts(zf)[0]
        assert st["applied"] == 1


def test_missing_char_and_multiple_revisions():
    """漏字 + 多条修订并存，批注条数与顺序正确。"""
    with tempfile.TemporaryDirectory() as td:
        src = _make_docx(os.path.join(td, "in.docx"),
                         ["通过字经济与实体经济融合，构建生太闭环"],)
        data, st = build_revised_docx(src, [
            {"find": "字经济", "replace": "数字经济"},
            {"find": "生太", "replace": "生态"},
        ])
        out = os.path.join(td, "o.docx")
        open(out, "wb").write(data)
        zf = zipfile.ZipFile(out)
        assert zf.testzip() is None
        texts = " ".join(_comment_texts(zf))
        assert len(_comments(zf)) == 2
        assert "数字经济" in texts and "生态" in texts
        assert st["applied"] == 2 and st["comments"] == 2


def test_all_occurrences_vs_once():
    """同一错字出现 2 次：默认全改（2 条批注），all_occurrences=False 只改 1 处。"""
    with tempfile.TemporaryDirectory() as td:
        src = _make_docx(os.path.join(td, "in.docx"), ["消废模式与消废场景"])
        d1, s1 = build_revised_docx(src, [{"find": "消废", "replace": "消费"}])
        d2, s2 = build_revised_docx(src, [{"find": "消废", "replace": "消费"}],
                                    all_occurrences=False)
        assert s1["applied"] == 2 and s1["comments"] == 2
        assert s2["applied"] == 1 and s2["comments"] == 1
        for d in (d1, d2):
            o = os.path.join(td, "o.docx")
            open(o, "wb").write(d)
            assert zipfile.ZipFile(o).testzip() is None


def test_no_match_reports_missed():
    """未命中必须显式回报，不能静默交付「看起来成功」的产物。"""
    with tempfile.TemporaryDirectory() as td:
        src = _make_docx(os.path.join(td, "in.docx"), ["这是一段正常文本"])
        data, st = build_revised_docx(src, [{"find": "根本不存在", "replace": "x"}])
        assert st["applied"] == 0
        assert st["missed"] == ["根本不存在"]
        # 仍应是健康可打开的 docx（无批注但格式合法）
        o = os.path.join(td, "o.docx")
        open(o, "wb").write(data)
        zf = zipfile.ZipFile(o)
        assert zf.testzip() is None


def test_table_paragraph_supported():
    """表格内段落同样能定位并加批注。"""
    with tempfile.TemporaryDirectory() as td:
        from docx import Document
        d = Document()
        t = d.add_table(rows=1, cols=1)
        t.cell(0, 0).text = "消废升级"
        src = os.path.join(td, "in.docx")
        d.save(src)
        data, st = build_revised_docx(src, [{"find": "消废", "replace": "消费"}])
        assert st["applied"] == 1 and st["comments"] == 1
        o = os.path.join(td, "o.docx")
        open(o, "wb").write(data)
        zf = zipfile.ZipFile(o)
        assert zf.testzip() is None
        assert "word/comments.xml" in zf.namelist()


def test_no_duplicate_zip_entries():
    """fresh-write 打包：zip 内不得出现重复条目（历史模型手写的主要坑）。"""
    with tempfile.TemporaryDirectory() as td:
        src = _make_docx(os.path.join(td, "in.docx"), ["消废与模市"])
        data, _ = build_revised_docx(src, [
            {"find": "消废", "replace": "消费"},
            {"find": "模市", "replace": "模式"},
        ])
        o = os.path.join(td, "o.docx")
        open(o, "wb").write(data)
        zf = zipfile.ZipFile(o)
        names = zf.namelist()
        assert len(names) == len(set(names)), f"存在重复条目: {names}"
        assert zf.testzip() is None


def test_original_format_preserved():
    """修订不该破坏原文格式：加粗与字号应保留。"""
    with tempfile.TemporaryDirectory() as td:
        src = _make_docx(os.path.join(td, "in.docx"), ["消废升级"], bold=True)
        data, _ = build_revised_docx(src, [{"find": "消废", "replace": "消费"}])
        o = os.path.join(td, "o.docx")
        open(o, "wb").write(data)
        root = ET.fromstring(zipfile.ZipFile(o).read("word/document.xml"))
        found = False
        for r in root.iter(W + "r"):
            txt = "".join(t.text or "" for t in r.iter(W + "t"))
            if txt == "消费":
                rPr = r.find(W + "rPr")
                assert rPr is not None
                assert rPr.find(W + "b") is not None, "加粗丢失"
                sz = rPr.find(W + "sz")
                assert sz is not None and sz.get(W + "val") == "24", "字号丢失"
                found = True
        assert found


def test_errors():
    """输入校验：源文件不存在 / 空修订列表应抛明确异常。"""
    with pytest.raises(DocxRevisionError):
        build_revised_docx("/nonexistent/x.docx", [{"find": "a", "replace": "b"}])
    with tempfile.TemporaryDirectory() as td:
        src = _make_docx(os.path.join(td, "in.docx"), ["ok"])
        with pytest.raises(DocxRevisionError):
            build_revised_docx(src, [])
        with pytest.raises(DocxRevisionError):
            build_revised_docx(src, [{"replace": "缺 find"}])
